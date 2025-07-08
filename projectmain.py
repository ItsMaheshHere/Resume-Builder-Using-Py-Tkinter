import tkinter as tk
from tkinter import ttk
top = tk.Tk()
from PIL import Image, ImageTk
import pygame
import json
from docx import Document
import os
top.configure(bg="Black")
top.geometry("1000x700")
top.title("   Resume Builder")

page1=tk.Frame(top,bg="black")
page2=tk.Frame(top,bg="black")
page3=tk.Frame(top,bg="black")
page4=tk.Frame(top,bg="black")
page5=tk.Frame(top,bg="black")

page1.place(x=0, y=0, relwidth=1, relheight=1)
page2.place(x=0, y=0, relwidth=1, relheight=1)
page3.place(x=0, y=0, relwidth=1, relheight=1)
page4.place(x=0, y=0, relwidth=1, relheight=1)
page5.place(x=0, y=0, relwidth=1, relheight=1)

page1.tkraise()

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

pygame.mixer.init()
click_sound = pygame.mixer.Sound(resource_path("click-sfx.mp3"))
click_sound.set_volume(0.8)
def play_click_sound():
    click_sound.play()

def handle_enter_key_page2(event):
    if var1.get() == 1:
        page3.tkraise()
    elif var2.get() == 1:
        page4.tkraise()

def create_resume_docx(user_details):
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    name = doc.add_paragraph()
    name_run = name.add_run(user_details["name"].upper())
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.space_after = Pt(8)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{user_details['mobile']} | {user_details['email']}"
    if user_details['linkedin']:
        contact_text += f" | {user_details['linkedin']}"
    contact_run = contact.add_run(contact_text)
    contact_run.font.size = Pt(11)
    contact.space_after = Pt(4)
    
    location = doc.add_paragraph()
    location.alignment = WD_ALIGN_PARAGRAPH.CENTER
    location_text = f"{user_details['city']}, {user_details['state']}"
    if user_details['website']:
        location_text += f" | {user_details['website']}"
    location_run = location.add_run(location_text)
    location_run.font.size = Pt(11)
    location.space_after = Pt(12)
    
    table = doc.add_table(rows=1, cols=2)
    table.allow_autofit = True
    
    left_cell = table.cell(0, 0)
    left_cell.width = Inches(3.5)
    
    edu_heading = left_cell.add_paragraph()
    edu_run = edu_heading.add_run('EDUCATION')
    edu_run.bold = True
    edu_run.font.size = Pt(12)
    edu_heading.space_after = Pt(6)
    
    edu = left_cell.add_paragraph()
    school_run = edu.add_run(f"{user_details['school']}\n")
    school_run.bold = True
    school_run.font.size = Pt(11)
    degree_run = edu.add_run(f"{user_details['degree']} in {user_details['field']}\n")
    degree_run.font.size = Pt(11)
    date_run = edu.add_run(f"{user_details['start_date']} - {user_details['end_date']}")
    date_run.font.size = Pt(10)
    date_run.italic = True
    edu.space_after = Pt(12)
    
    if user_details['skills']:
        skills_heading = left_cell.add_paragraph()
        skills_run = skills_heading.add_run('SKILLS')
        skills_run.bold = True
        skills_run.font.size = Pt(12)
        skills_heading.space_after = Pt(6)
        
        skills = left_cell.add_paragraph()
        for skill in user_details['skills']:
            skill_run = skills.add_run('• ' + skill + '\n')
            skill_run.font.size = Pt(11)
        skills.space_after = Pt(12)
    
    if user_details['languages']:
        lang_heading = left_cell.add_paragraph()
        lang_run = lang_heading.add_run('LANGUAGES')
        lang_run.bold = True
        lang_run.font.size = Pt(12)
        lang_heading.space_after = Pt(6)
        
        languages = left_cell.add_paragraph()
        lang_run = languages.add_run(user_details['languages'])
        lang_run.font.size = Pt(11)
        languages.space_after = Pt(12)
    
    right_cell = table.cell(0, 1)
    right_cell.width = Inches(4)
    
    if user_details['summary']:
        summary_heading = right_cell.add_paragraph()
        summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
        summary_run.bold = True
        summary_run.font.size = Pt(12)
        summary_heading.space_after = Pt(6)
        
        summary = right_cell.add_paragraph()
        summary_run = summary.add_run(user_details['summary'])
        summary_run.font.size = Pt(11)
        summary.space_after = Pt(12)
    
    if user_details['experience']:
        exp_heading = right_cell.add_paragraph()
        exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
        exp_run.bold = True
        exp_run.font.size = Pt(12)
        exp_heading.space_after = Pt(6)
        
        experiences = user_details['experience'].split('\n\n')
        for exp in experiences:
            experience = right_cell.add_paragraph()
            exp_run = experience.add_run(exp)
            exp_run.font.size = Pt(11)
            experience.space_after = Pt(8)
    
    for row in table.rows:
        row.height = Inches(8)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    for cell in table.columns[0].cells:
        cell.width = Inches(3.5)
    for cell in table.columns[1].cells:
        cell.width = Inches(4)
    return doc

def create_pdf_resume(user_details, is_professional=False):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    # Create PDF
    output_filename = f"resume_{user_details['name'].replace(' ', '_')}.pdf"
    doc = SimpleDocTemplate(output_filename, pagesize=letter,rightMargin=0.5*inch, leftMargin=0.5*inch,topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    
    styles.add(ParagraphStyle(name='CustomTitle',parent=styles['Heading1'],fontSize=24,spaceAfter=8,alignment=TA_CENTER))
    
    styles.add(ParagraphStyle(name='ContactInfo',parent=styles['Normal'],fontSize=11,spaceAfter=4,alignment=TA_CENTER))
    
    styles.add(ParagraphStyle(name='SectionHeading',parent=styles['Heading2'],fontSize=12,spaceBefore=12,spaceAfter=6,textColor=colors.black,bold=True))
    
    styles.add(ParagraphStyle(name='NormalText',parent=styles['Normal'],fontSize=11,spaceAfter=8))
    content = []
    content.append(Paragraph(user_details["name"].upper(), styles['CustomTitle']))
    
    contact_info = []
    contact_parts = [user_details['mobile'], user_details['email']]
    if user_details['linkedin']:
        contact_parts.append(user_details['linkedin'])
    contact_info.append(Paragraph(' | '.join(contact_parts), styles['ContactInfo']))
    
    location_parts = [f"{user_details['city']}, {user_details['state']}"]
    if user_details['website']:
        location_parts.append(user_details['website'])
    contact_info.append(Paragraph(' | '.join(location_parts), styles['ContactInfo']))
    
    content.extend(contact_info)
    content.append(Spacer(1, 12))
    
    left_column = []
    right_column = []
    
    left_column.append(Paragraph('EDUCATION', styles['SectionHeading']))
    education_text = [
        f"<b>{user_details['school']}</b>",
        f"{user_details['degree']} in {user_details['field']}",
        f"<i>{user_details['start_date']} - {user_details['end_date']}</i>"]

    left_column.append(Paragraph('<br/>'.join(education_text), styles['NormalText']))
    
    if user_details['skills']:
        left_column.append(Paragraph('SKILLS', styles['SectionHeading']))
        skills_text = ''.join([f"• {skill}<br/>" for skill in user_details['skills']])
        left_column.append(Paragraph(skills_text, styles['NormalText']))
    
    if user_details['languages']:
        left_column.append(Paragraph('LANGUAGES', styles['SectionHeading']))
        left_column.append(Paragraph(user_details['languages'], styles['NormalText']))
    
    if user_details['summary']:
        right_column.append(Paragraph('PROFESSIONAL SUMMARY', styles['SectionHeading']))
        right_column.append(Paragraph(user_details['summary'], styles['NormalText']))
    
    if user_details['experience']:
        right_column.append(Paragraph('PROFESSIONAL EXPERIENCE', styles['SectionHeading']))
        experiences = user_details['experience'].split('\n\n')
        for exp in experiences:
            right_column.append(Paragraph(exp, styles['NormalText']))
    
    table_data = [[left_column, right_column]]
    col_widths = [doc.width/2.5, doc.width - (doc.width/2.5)]
    table = Table(table_data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 20),
    ]))
    content.append(table)
    doc.build(content)
    return output_filename

def handle_enter_key_page5(event):
    if format_var1.get() == 1: 
        
        experience_text = None
        is_professional = False
            
        if var1.get() == 1: 
            for widget in page3.winfo_children():
                if isinstance(widget, tk.Text):
                    experience_text = widget.get("1.0", tk.END).strip()
                    break
        else: 
            experience_text = prof_experience_text.get("1.0", tk.END).strip()
            is_professional = True
                    
        user_details = {
            "name": entry.get(),
            "mobile": en1.get(),
            "linkedin": en2.get(),
            "email": en3.get(),
            "city": en4.get(),
            "state": en5.get(),
            "website": en6.get() if combo.get() == "Yes" else "",
            "school": en7.get(),
            "degree": en8.get(),
            "field": en9.get(),
            "start_date": en10.get(),
            "end_date": en11.get(),
            "summary": summary_text.get("1.0", tk.END).strip(),
            "languages": languages_text.get("1.0", tk.END).strip(),
            "skills": app.selected_skills,
            "experience": experience_text
        }
        output_filename = create_pdf_resume(user_details, is_professional)
        print(f"Resume successfully saved as {output_filename}")
            
        success_label = tk.Label(page5, text=f"Resume saved as {output_filename}", bg="Black", fg="Green", font=("Arial", 12))
        success_label.place(x=350, y=250)
            
    elif format_var2.get() == 1:  
        
        experience_text = None
        if var1.get() == 1: 
            for widget in page3.winfo_children():
                if isinstance(widget, tk.Text):
                    experience_text = widget.get("1.0", tk.END).strip()
                    break
        else:  
            experience_text = prof_experience_text.get("1.0", tk.END).strip()
                    
        user_details = {
            "name": entry.get(),
            "mobile": en1.get(),
            "linkedin": en2.get(),
            "email": en3.get(),
            "city": en4.get(),
            "state": en5.get(),
            "website": en6.get() if combo.get() == "Yes" else "",
            "school": en7.get(),
            "degree": en8.get(),
            "field": en9.get(),
            "start_date": en10.get(),
            "end_date": en11.get(),
            "summary": summary_text.get("1.0", tk.END).strip(),
            "languages": languages_text.get("1.0", tk.END).strip(),
            "skills": app.selected_skills,
            "experience": experience_text
        }
            
        print("Creating new resume document...")
        doc = create_resume_docx(user_details)
            
        output_filename = f"resume_{user_details['name'].replace(' ', '_')}.docx"
        print(f"Saving resume as: {output_filename}")
        doc.save(output_filename)
        print(f"Resume successfully saved as {output_filename}")
            
        success_label = tk.Label(page5, text=f"Resume saved as {output_filename}", bg="Black", fg="Green", font=("Arial", 12))
        success_label.place(x=350, y=250)
            
def bind_enter_key():
    top.unbind("<Return>")
    if page2.winfo_viewable():
        top.bind("<Return>", handle_enter_key_page2)
    elif page5.winfo_viewable():
        top.bind("<Return>", handle_enter_key_page5)

def next_btn():
    page2.tkraise()
    play_click_sound()
    bind_enter_key()

def back_btn():
    page1.tkraise()
    play_click_sound()
    bind_enter_key()

def next3_btn():
    page5.tkraise()
    play_click_sound()
    top.unbind("<Return>")
    top.bind("<Return>", handle_enter_key_page5)

def back2_btn():
    page2.tkraise()
    play_click_sound()
    bind_enter_key()

def back5_btn():
    page3.tkraise()
    play_click_sound()
    top.unbind("<Return>")
    top.bind("<Return>", handle_enter_key)

next__btn = tk.Button(page1, text="Next",fg="White",bg="Black", command=next_btn)
next__btn.place(x=1300,y=580)
back__btn=tk.Button(page2,text="Back",fg="White",bg="Black",command=back_btn)
back__btn.place(x=1260,y=580)

back_2_btn=tk.Button(page4,text="Back",fg="White",bg="Black",command=back2_btn)
back_2_btn.place(x=1260,y=580)

next_3_btn= tk.Button(page3, text="Next",fg="White",bg="Black", command=next3_btn)
next_3_btn.place(x=1300,y=580)
back_3_btn=tk.Button(page3,text="Back",fg="White",bg="Black",command=back2_btn)
back_3_btn.place(x=1260,y=580)

back_3_btn=tk.Button(page5,text="Back",fg="White",bg="Black",command=back5_btn)
back_3_btn.place(x=1260,y=580)

next_4_btn=tk.Button(page4, text="Next",fg="White",bg="Black", command=next3_btn)
next_4_btn.place(x=1300,y=580)

style = ttk.Style()
style.theme_use("default")
style.configure("CustomCombobox.TCombobox",fg="white",bg="black",fieldbackground="black",arrowcolor="orange")

def yesorno(event):
    if combo.get() == "Yes":
        en0.place(x=180, y=370)
        line8.place(x=180,y=390)
    else:
        en0.place_forget()       
        line8.place_forget()

class SkillSelector:
    def __init__(self, master):
        
        with open(resource_path("skills.json"),"r") as file:
            data=json.load(file)
            self.all_skills=data["skills"]
        self.selected_skills = []

        self.skill_entry = tk.Entry(page1, width=20,bg="Black",fg="White",insertbackground="white",font=('Arial', 12))
        self.skill_entry.place(x=180,y=440)
        
        self.suggestion_listbox = tk.Listbox(page1,width=30,bg="Black",fg="White",font=('Arial', 12),height=5,selectmode=tk.SINGLE)
        self.suggestion_listbox.place(x=180,y=480)
        
        self.skill_entry.bind('<KeyRelease>', self.update_suggestions)
        self.skill_entry.bind('<Return>', self.add_selected_skill)
        self.suggestion_listbox.bind('<Button-1>', self.on_select_suggestion)
        
    def update_suggestions(self, event=None):
        search_term = self.skill_entry.get().lower()
        self.suggestion_listbox.delete(0, tk.END)
        
        if search_term:
            matches = [skill for skill in self.all_skills 
                      if skill.lower().startswith(search_term)]
            for skill in matches:
                self.suggestion_listbox.insert(tk.END, skill)
    
    def on_select_suggestion(self, event=None):
        if self.suggestion_listbox.curselection():
            selected = self.suggestion_listbox.get(self.suggestion_listbox.curselection())
            self.skill_entry.delete(0, tk.END)
            self.skill_entry.insert(0, selected)
    
    def add_selected_skill(self, event=None):
        skill = self.skill_entry.get().strip()
        if skill and skill not in self.selected_skills:
            self.selected_skills.append(skill)
            self.display_selected_skills()
            self.skill_entry.delete(0, tk.END)
    
    def display_selected_skills(self):
        for widget in page1.winfo_children():
            if isinstance(widget, tk.Label) and widget.cget("bg") == "navy":
                widget.destroy()
        x_pos = 180 
        y_pos = 600 
        max_width = 180 + self.suggestion_listbox.winfo_width()
        
        for skill in self.selected_skills:
            skill_label = tk.Label(page1, text=skill, bg="navy", fg="white", padx=10, pady=5, font=('Arial', 10))
            
            if x_pos + skill_label.winfo_reqwidth() > max_width:
                x_pos = 180 
                y_pos += 40 
            
            skill_label.place(x=x_pos, y=y_pos)
            x_pos += skill_label.winfo_reqwidth() + 20

def get_summary():
    summary=summary_text.get("1.0", tk.END).strip()
    print("Your resume summary: ", summary)
    return summary

def get_languages():
    languages=languages_text.get("1.0", tk.END).strip()
    print("Your known languages: ", languages)
    return languages

def get_education():
    education={"school": en7.get(),"degree": en8.get(),"field": en9.get(),"start_date": en10.get(),"end_date": en11.get()}
    print("Your education details: ", education)
    return education

def get_experience():
    experience=text.get("1.0", tk.END)
    experience=experience.rstrip('\n')
    print("Your experience:")
    print(experience)
    return experience

lbl=tk.Label(page1,text="Welcome To Resume Builder",font=("Comic Sans MS", 22, "bold"),bg="Black",fg="White")
lbl.pack()
canvas=tk.Canvas(page1,width=600,height=5,bg="Black",highlightthickness=0)
canvas.create_line(50,2,550,2,fill="Yellow",width=2)
canvas.pack(pady=10)

lbl=tk.Label(page1,text="Enter Your Details :-",font=("Courier New",15, "bold"),bg="#FF5733",fg="White")
lbl.place(x=10,y=80)


#1)Details
lbl1=tk.Label(page1,text="Your Full Name    :-",bg="Black",fg="White",font=(11))
lbl1.place(x=30,y=120)
def name(event):
    name = entry.get()
    print("Your Name: ",name)
entry = tk.Entry(page1, bg="black", fg="white", bd=0, insertbackground="white",width=17,font=(16))
entry.place(x=180, y=119)
entry.bind("<Return>", name)
line = tk.Frame(page1, bg="white", height=2, width=150)
line.place(x=180,y=139)

lbl2=tk.Label(page1,text="Enter Mobile No. :-",bg="Black",fg="White",font=(11))
lbl2.place(x=30,y=160)
def mobile_no(event):
    mobileno = en1.get()
    print("Mobile No.:-",mobileno)
en1=tk.Entry(page1, bg="black", fg="white", bd=0, insertbackground="white",width=17,font=(16))
en1.place(x=180, y=156)
en1.bind("<Return>",mobile_no)
line = tk.Frame(page1, bg="white", height=2, width=150)
line.place(x=180,y=176)

img_raw = Image.open(resource_path("linkedin_Logo.png"))
img_resized = img_raw.resize((27, 27))
img = ImageTk.PhotoImage(img_resized)
label = tk.Label(page1, image=img,bg="black")
label.image = img  
label.place(x=27, y=195)

lbl3=tk.Label(page1,text="Profile Link   :-",bg="Black",fg="White",font=(11))
lbl3.place(x=60,y=200)
def linkedin_link(event):
    linkedin=en2.get()
    print("Your Link. Profile Link:- ", linkedin)
en2=tk.Entry(page1,bg="Black",fg="White",bd=0,insertbackground="White",width=30,font=(16))
en2.place(x=180,y=200)
en2.bind("<Return>",linkedin_link)
line = tk.Frame(page1, bg="white", height=2, width=265)
line.place(x=180,y=220)

lbl4 = tk.Label(page1,text="Email Address    :-",bg="Black",fg="White",font=(11))
lbl4.place(x=30,y=240)
def email_add(event):
    Email=en3.get()
    print("Your mail id:- ",Email)
en3 = tk.Entry(page1,bg="Black",fg="White",bd=0,insertbackground="White",width=30,font=(16))
en3.place(x=180,y=240)
en3.bind("<Return>",email_add)
line = tk.Frame(page1, bg="white", height=2, width=265)
line.place(x=180,y=260)

lbl5 = tk.Label(page1,text="Location :-",bg="Black",fg="White",font=(11))
lbl5.place(x=30,y=270)
lbl = tk.Label(page1,text="City -",bg="Black",fg="White",font=(11))
lbl.place(x=120,y=274)
def Location_city(event):
    city=en4.get()
    print("Your City is: ",city)
en4 =  tk.Entry(page1,bg="Black",fg="White",bd=0,insertbackground="White",width=13,font=(16))
en4.place(x=180,y=270)
en4.bind("<Return>",Location_city)
line = tk.Frame(page1, bg="white", height=2, width=110)
line.place(x=180,y=290)

lbl6=tk.Label(page1,text="State -",bg="Black",fg="White",font=(11))
lbl6.place(x=113,y=304)
def Location_state(event):
    state=en5.get()
    print("Your State is:- ",state)
en5 =  tk.Entry(page1,bg="Black",fg="White",bd=0,insertbackground="White",width=13,font=(16))
en5.place(x=180,y=300)
en5.bind("<Return>",Location_state)
line = tk.Frame(page1, bg="white", height=2, width=110)
line.place(x=180,y=320)

lbl7=tk.Label(page1,text="Personal Website:-",bg="Black",fg="White",font=(11))
lbl7.place(x=30,y=340)
combo = ttk.Combobox(page1, values=["Yes", "No"],style="CustomCombobox.TCombobox")
combo.bind("<<ComboboxSelected>>", yesorno)
combo.place(x=180,y=340)
def per_web(event):
    website = en6.get()
    print("Your website: ", website)
    return website
en6=tk.Entry(page1,bg="Black",fg="White",bd=0,insertbackground="White",width=30,font=(16))
en6.bind("<Return>", per_web)
line8 = tk.Frame(page1, bg="white", height=2, width=265)

lbl8=tk.Label(page1,text="Enter Your Skills :-",font=("Courier New",15, "bold"),bg="#FF5733",fg="White")
lbl8.place(x=10,y=400)
lbl9=tk.Label(page1,text="Enter here  :-",bg="Black",fg="White",font=(11))
lbl9.place(x=70,y=440)
app = SkillSelector(page1)

lbl10=tk.Label(page1,text="Write a summary of Your Resume :-",font=("Courier New",15, "bold"),bg="#FF5733",fg="White")
lbl10.place(x=550,y=80)
summary_text = tk.Text(page1, height=10, width=50,bg="Black",fg="White",font=("Arial", 14),insertbackground="White")
summary_text.place(x=640,y=120)
summary_text.bind("<Return>", lambda e: get_summary())

lbl11=tk.Label(page1,text="Your Education :-",font=("Courier New",15, "bold"),bg="#FF5733",fg="White")
lbl11.place(x=550,y=400)

lbl12=tk.Label(page1,text="School/University:-",bg="Black",fg="White",font=(11))
lbl12.place(x=630,y=435)
en7=tk.Entry(page1,bg="Black",fg="White",insertbackground="White",width=23,font=(16),bd=0)
en7.place(x=765,y=435)
line=tk.Frame(page1, bg="white", height=2, width=205)
line.place(x=765,y=455)

lbl13=tk.Label(page1,text="Degree :-",bg="Black",fg="White",font=(11))
lbl13.place(x=690,y=464)
en8=tk.Entry(page1,bg="Black",fg="White",insertbackground="White",width=23,font=(16),bd=0)
en8.place(x=765,y=465)
line=tk.Frame(page1, bg="white", height=2, width=205)
line.place(x=765,y=485)

en0=tk.Entry(page1,bg="Black",fg="White",insertbackground="White",width=23,font=(16),bd=0)

lbl14=tk.Label(page1,text="Field of study :-",bg="Black",fg="White",font=(11))
lbl14.place(x=651,y=493)
en9=tk.Entry(page1,bg="Black",fg="White",insertbackground="White",width=23,font=(16),bd=0)
en9.place(x=765,y=494)
line=tk.Frame(page1, bg="white", height=2, width=205)
line.place(x=765,y=513)

lbl15=tk.Label(page1,text="Start Date :-",bg="Black",fg="White",font=(11))
lbl15.place(x=670,y=527)
en10=tk.Entry(page1,bg="Black",fg="White",insertbackground="White",width=9,font=(16),bd=0)
en10.place(x=765,y=528)
line=tk.Frame(page1, bg="white", height=2, width=80)
line.place(x=765,y=548)

lbl16=tk.Label(page1,text="End Date :-",bg="Black",fg="White",font=(11))
lbl16.place(x=672,y=558)
en11=tk.Entry(page1,fg="White",bg="Black",insertbackground="White",width=9,font=(16),bd=0)
en11.place(x=765,y=558)
line=tk.Frame(page1, bg="white", height=2, width=80)
line.place(x=765,y=578)

lbl=tk.Label(page1,text="[DD/MM/YYYY]",bg="Black",fg="White",font=(200))
lbl.place(x=558,y=527)
lbl=tk.Label(page1,text="[DD/MM/YYYY]",bg="Black",fg="White",font=(200))
lbl.place(x=558,y=558)

en7.bind("<Return>", lambda e: get_education())
en8.bind("<Return>", lambda e: get_education())
en9.bind("<Return>", lambda e: get_education())
en10.bind("<Return>", lambda e: get_education())
en11.bind("<Return>", lambda e: get_education())

lbl01=tk.Label(page1,text="Languages:-",font=("Courier New",15, "bold"),bg="#FF5733",fg="White")
lbl01.place(x=1000,y=400)
languages_text = tk.Text(page1, height=5, width=20,bg="Black",fg="White",font=("Arial", 14),insertbackground="White")
languages_text.place(x=1020,y=440)
languages_text.bind("<Return>", lambda e: get_languages())

#page2
var1 = tk.IntVar()
var2 = tk.IntVar()

options_frame = tk.Frame(page2, bg="black")
options_frame.place(relx=0.5, rely=0.5, anchor="center")

title_label = tk.Label(options_frame, text="Select Resume Format", font=("Comic Sans MS", 22, "bold"), bg="Black", fg="White")
title_label.pack(pady=20)

entry_frame = tk.Frame(options_frame, bg="black")
entry_frame.pack(pady=20)
entry_label = tk.Label(entry_frame, text="Entry Level Resume", font=("Arial", 14), bg="Black", fg="White")
entry_label.pack(side="left", padx=10)
entry_checkbox = tk.Checkbutton(entry_frame, bg="Black", fg="White", selectcolor='black', variable=var1)
entry_checkbox.pack(side="left")

prof_frame = tk.Frame(options_frame, bg="black")
prof_frame.pack(pady=20)
prof_label = tk.Label(prof_frame, text="Professional Level Resume", font=("Arial", 14), bg="Black", fg="White")
prof_label.pack(side="left", padx=10)
prof_checkbox = tk.Checkbutton(prof_frame, bg="Black", fg="White", selectcolor='black', variable=var2)
prof_checkbox.pack(side="left")

instruction_label = tk.Label(options_frame, text="Press Enter after selection", bg="Black", fg="orange", font=("Arial", 12))
instruction_label.pack(pady=20)

def handle_enter_key(event):
    if var1.get() == 1:
        page3.tkraise()
    elif var2.get() == 1:
        page4.tkraise()

top.bind("<Return>", handle_enter_key)

lbl=tk.Label(page3,text="For this Format you have to enter one more detail :-",font=("Comic Sans MS", 22, "bold"),bg="Black",fg="orange")
lbl.pack()
lbl=tk.Label(page3,text="Write the Experience Whatever you have:-",bg="Black",fg="White",font=(11))
lbl.place(x=30,y=70)
text=tk.Text(page3,height=20, width=70,bg="Black",fg="White",font=("Arial", 14),insertbackground="White")
text.place(x=40,y=100)

text.bind("<Return>", lambda e: get_experience())

#page5
lbl=tk.Label(page5,text="Select Resume Format",font=("Comic Sans MS", 22, "bold"),bg="Black",fg="orange")
lbl.pack()

format_var1 = tk.IntVar()
format_var2 = tk.IntVar()

lbl_format1 = tk.Label(page5, text="PDF Format", bg="Black", fg="White", font=(11))
lbl_format1.place(x=400, y=100)
cbtn_format1 = tk.Checkbutton(page5, text="", bg="Black", fg="White", selectcolor='black', variable=format_var1)
cbtn_format1.place(x=580, y=100)

lbl_format2 = tk.Label(page5, text="Word Format", bg="Black", fg="White", font=(11))
lbl_format2.place(x=400, y=150)
cbtn_format2 = tk.Checkbutton(page5, text="", bg="Black", fg="White", selectcolor='black', variable=format_var2)
cbtn_format2.place(x=580, y=150)

lbl_instruction = tk.Label(page5, text="Select one format and press Enter to continue", bg="Black", fg="orange", font=("Bold", 14))
lbl_instruction.place(x=350, y=200)

lbl=tk.Label(page4, text="Enter Your Professional Experience Details:", font=("Comic Sans MS", 22, "bold"), bg="Black", fg="orange")
lbl.pack()

lbl_instructions = tk.Label(page4, text="Include your work experience, job titles, companies, and key achievements:", bg="Black", fg="White", font=(11))
lbl_instructions.place(x=30, y=70)

prof_experience_text = tk.Text(page4, height=20, width=70, bg="Black", fg="White", font=("Arial", 14), insertbackground="White")
prof_experience_text.place(x=40, y=100)

def get_professional_experience(event=None):
    experience = prof_experience_text.get("1.0", tk.END).strip()
    print("Professional Experience:")
    print(experience)
    return experience

prof_experience_text.bind("<Return>", get_professional_experience)

top.mainloop()